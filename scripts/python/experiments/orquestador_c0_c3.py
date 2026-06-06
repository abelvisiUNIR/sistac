"""
orquestador_c0_c3.py — Orquestador del experimento factorial C0-C3

Ejecuta las cuatro condiciones experimentales y consolida los resultados
en tablas Word (.docx) y CSV para insertar en paper/SISTAC_TFE.docx.

Autores: David I. Madrid Oyanadel + Mario A. Belvisi Lescano

CÓMO EJECUTAR:
    python scripts/python/experiments/orquestador_c0_c3.py

OUTPUTS (en paper/tables/):
    tab_resultados_h1.docx   — Tabla de tiempos (H1)
    tab_resultados_h1.csv
    tab_resultados_h2.docx   — Tabla F1/AUC-ROC (H2)
    tab_resultados_h2.csv
    tab_resultados_h3.docx   — Tabla DIR/SPD (H3)
    tab_resultados_h3.csv
    tab_comparativa.docx     — Comparativa integrada C0-C3

PREREQUISITOS:
    1. ChromaDB o Azure AI Search configurado (ver .env.example)
    2. Pipeline indexado: SistacRAGPipeline.index() ejecutado
    3. Corpus en data/raw/cvs/ y data/raw/job_descriptions/
    4. Gold Standard en data/raw/gold_standard/ground_truth.csv
    5. Tiempos C0 en data/raw/gold_standard/c0_times.csv

FORMATO ground_truth.csv:
    cv_id,jd_id,expected_label,expected_score,group_gender,group_age
    CV_001,JD_001,APTO,90,F,25-34

FORMATO c0_times.csv:
    cv_id,jd_id,time_seconds,decision,evaluator_id
    CV_001,JD_001,420.5,APTO,EVAL_01
"""

from __future__ import annotations

import csv
import json
import random
import sys
from pathlib import Path

import numpy as np

# INV-14: semilla global
random.seed(42)
np.random.seed(42)

# INV-16: rutas via PROJECT_ROOT
SCRIPTS_DIR = Path(__file__).parent.parent
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from config import (
    TABLES_DIR,
    EVAL_SETS,
    GOLD_STANDARD_DIR,
    CVS_RAW,
    JOB_DESCRIPTIONS,
    ensure_dirs,
)
from evaluation.efficacy_metrics import efficacy_report
from evaluation.efficiency_metrics import efficiency_report
from evaluation.fairness_metrics import fairness_report
from rag.pipeline import SistacRAGPipeline


# ── C0: Baseline manual ───────────────────────────────────────────────────────

def run_c0_baseline() -> list[dict]:
    """
    Carga los tiempos y decisiones del screening manual (C0 — Gold Standard).

    Los evaluadores RRHH registraron sus tiempos en c0_times.csv.
    Este archivo debe existir antes de ejecutar el orquestador.

    Returns:
        Lista de dicts con: cv_id, jd_id, time_seconds, decision, config="c0"
    """
    c0_path = GOLD_STANDARD_DIR / "c0_times.csv"
    if not c0_path.exists():
        raise FileNotFoundError(
            f"Archivo de tiempos C0 no encontrado: {c0_path}\n"
            "Registrar los tiempos del screening manual antes de ejecutar."
        )

    results = []
    with open(c0_path, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            results.append({
                "cv_id":        row["cv_id"],
                "jd_id":        row["jd_id"],
                "time_seconds": float(row["time_seconds"]),
                "decision":     row["decision"],
                "config":       "c0",
            })

    print(f"  [C0] {len(results)} evaluaciones cargadas desde {c0_path.name}")
    return results


# ── C1, C2, C3: Sistemas automáticos ─────────────────────────────────────────

CACHE_FILE = Path("/app/data/eval_cache.json")

def load_cache() -> dict:
    if CACHE_FILE.exists():
        try:
            with open(CACHE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            print(f"  [WARN] Error cargando caché de evaluaciones: {e}")
    return {}

def save_cache(cache: dict) -> None:
    try:
        CACHE_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(CACHE_FILE, "w", encoding="utf-8") as f:
            json.dump(cache, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"  [WARN] Error guardando caché de evaluaciones: {e}")


def _run_pipeline(config: str, eval_pairs: list[dict]) -> list[dict]:
    """
    Ejecuta el pipeline RAG para una configuración experimental.

    Args:
        config:     "c1", "c2" o "c3".
        eval_pairs: Lista de dicts {cv_id, jd_id, cv_text, jd_text}.

    Returns:
        Lista de dicts con resultados de evaluate() más expected_label y group.
    """
    pipeline = SistacRAGPipeline(config=config)
    results = []

    cache = load_cache()

    for pair in eval_pairs:
        cache_key = f"{config}_{pair['cv_id']}_{pair['jd_id']}"

        # Si ya existe en la caché local, recuperamos el resultado
        if cache_key in cache:
            result = cache[cache_key]
            # Nos aseguramos de inyectar la metadata actual del Gold Standard
            result["expected_label"] = pair.get("expected_label", "")
            result["expected_score"] = pair.get("expected_score", None)
            result["group_gender"]   = pair.get("group_gender", "")
            result["group_age"]      = pair.get("group_age", "")
            results.append(result)
            continue

        print(f"  [{config.upper()}] Evaluando {pair['cv_id']} ↔ {pair['jd_id']}...")
        result = pipeline.evaluate(
            cv_id=pair["cv_id"],
            cv_text=pair["cv_text"],
            jd_id=pair["jd_id"],
            jd_text=pair["jd_text"],
        )
        # Agregar metadata del Gold Standard para cálculo de métricas
        result["expected_label"] = pair.get("expected_label", "")
        result["expected_score"] = pair.get("expected_score", None)
        result["group_gender"]   = pair.get("group_gender", "")
        result["group_age"]      = pair.get("group_age", "")
        
        # Guardamos en caché y en disco de inmediato
        cache[cache_key] = result
        save_cache(cache)
        
        results.append(result)

    print(f"  [{config.upper()}] {len(results)} evaluaciones completadas.")
    return results



def run_c1_llm_pure(eval_pairs: list[dict]) -> list[dict]:
    """C1: LLM puro sin RAG."""
    return _run_pipeline("c1", eval_pairs)


def run_c2_llm_rag(eval_pairs: list[dict]) -> list[dict]:
    """C2: LLM + RAG (Azure AI Search)."""
    return _run_pipeline("c2", eval_pairs)


def run_c3_llm_rag_pii(eval_pairs: list[dict]) -> list[dict]:
    """C3: PII anonimización + LLM + RAG."""
    return _run_pipeline("c3", eval_pairs)


# ── Carga del dataset de evaluación ──────────────────────────────────────────

def load_eval_pairs() -> list[dict]:
    """
    Carga los pares (CV, JD) para el experimento desde el Gold Standard.

    Lee ground_truth.csv para obtener los pares a evaluar, luego carga
    el texto de cada CV y JD desde data/raw/.

    Returns:
        Lista de dicts con: cv_id, jd_id, cv_text, jd_text,
                            expected_label, expected_score,
                            group_gender, group_age
    """
    gt_path = GOLD_STANDARD_DIR / "ground_truth.csv"
    if not gt_path.exists():
        raise FileNotFoundError(
            f"Ground truth no encontrado: {gt_path}\n"
            "Generar el Gold Standard antes de ejecutar el orquestador."
        )

    pairs = []
    with open(gt_path, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            cv_path = CVS_RAW / f"{row['cv_id']}.txt"
            jd_path = JOB_DESCRIPTIONS / f"{row['jd_id']}.txt"

            if not cv_path.exists():
                print(f"  [WARN] CV no encontrado: {cv_path} — omitiendo")
                continue
            if not jd_path.exists():
                print(f"  [WARN] JD no encontrada: {jd_path} — omitiendo")
                continue

            pairs.append({
                "cv_id":          row["cv_id"],
                "jd_id":          row["jd_id"],
                "cv_text":        cv_path.read_text(encoding="utf-8"),
                "jd_text":        jd_path.read_text(encoding="utf-8"),
                "expected_label": row.get("expected_label", ""),
                "expected_score": row.get("expected_score", None),
                "group_gender":   row.get("group_gender", ""),
                "group_age":      row.get("group_age", ""),
            })

    print(f"Dataset cargado: {len(pairs)} pares CV-JD")
    return pairs


# ── Cálculo de métricas ───────────────────────────────────────────────────────

def compute_h1_metrics(c0_results: list[dict], cx_results: list[dict], config_name: str) -> dict:
    """H1 — Eficiencia: compara T_cand de C0 vs Cx (Mann-Whitney U)."""
    times_c0 = [r["time_seconds"] for r in c0_results]
    times_cx = [r["time_seconds"] for r in cx_results]
    return efficiency_report(times_baseline=times_c0, times_system=times_cx, config_name=config_name)



def compute_h2_metrics(results: list[dict], config_name: str) -> dict:
    """H2 — Eficacia: F1 macro y AUC-ROC del sistema vs Gold Standard."""
    y_true = np.array([1 if r["expected_label"] == "APTO" else 0 for r in results])
    y_pred = np.array([1 if r["decision"] == "APTO" else 0 for r in results])
    y_score = np.array([r["score"] / 100.0 if r["score"] is not None else 0.5 for r in results])
    res = efficacy_report(y_true=y_true, y_pred=y_pred, y_score=y_score, config_name=config_name)
    return {
        "f1": res.get("F1_macro", 0.0),
        "auc_roc": res.get("AUC_ROC", 0.0),
        "auc_ci_low": res.get("AUC_ROC_CI_lower", 0.0),
        "auc_ci_high": res.get("AUC_ROC_CI_upper", 0.0),
        "h2_accepted": res.get("H2_accepted", False)
    }


def compute_h3_metrics(results: list[dict], config_name: str, group_attr: str = "group_gender") -> dict:
    """H3 — Equidad: DIR y SPD entre grupos protegidos."""
    y_pred = np.array([1 if r["decision"] == "APTO" else 0 for r in results])
    group = np.array([r.get(group_attr, "desconocido") for r in results])
    
    # Privilegiado: "M" para género, "23-35" para edad
    priv_group = "M" if group_attr == "group_gender" else "23-35"
    
    res = fairness_report(
        y_pred=y_pred,
        group=group,
        config_name=config_name,
        privileged_group=priv_group,
        positive_label=1
    )
    
    return {
        "dir": res.get("DIR", 0.0),
        "spd": res.get("SPD", 0.0),
        "selection_rate_privileged": res.get("selection_rate_privileged", 0.0),
        "selection_rate_minority": res.get("selection_rate_minority", 0.0)
    }


# ── Guardar resultados ────────────────────────────────────────────────────────

def save_csv(data: list[dict], filename: str) -> None:
    """Guarda resultados como CSV en paper/tables/."""
    ensure_dirs()
    path = TABLES_DIR / filename
    if not data:
        return
    with open(path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=data[0].keys())
        writer.writeheader()
        writer.writerows(data)
    print(f"  → CSV: {path}")


def save_word_table(
    headers: list[str],
    rows: list[list[str]],
    caption: str,
    filename: str,
    note: str = "",
) -> None:
    """
    Guarda una tabla de resultados como .docx en paper/tables/.

    El archivo generado puede insertarse en SISTAC_TFE.docx copiando
    la tabla (Ctrl+A en el .docx → pegar en SISTAC_TFE.docx).
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  [WARN] python-docx no instalado. Solo se generará el CSV.")
        return

    ensure_dirs()
    doc = DocxDocument()

    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.bold = True
    run.font.size = Pt(11)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            row_cells[c_idx].paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )

    if note:
        p_note = doc.add_paragraph()
        run_note = p_note.add_run(f"Nota. {note}")
        run_note.italic = True
        run_note.font.size = Pt(10)

    path = TABLES_DIR / filename
    doc.save(str(path))
    print(f"  → DOCX: {path}")


# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    print("=" * 60)
    print("SISTAC — Experimento Factorial C0-C3")
    print("=" * 60)

    ensure_dirs()

    # 1. Cargar dataset de evaluación
    print("\n[1/6] Cargando dataset de evaluación...")
    try:
        eval_pairs = load_eval_pairs()
    except FileNotFoundError as e:
        print(f"\n[ERROR] {e}")
        print("\nEjecutando en modo TEMPLATE (sin datos reales).")
        _generate_template_tables()
        return

    # 2. Ejecutar C0 (baseline manual)
    print("\n[2/6] C0 — Screening manual (cargando tiempos)...")
    try:
        c0_results = run_c0_baseline()
    except FileNotFoundError as e:
        print(f"  [WARN] {e}\nC0 omitido — H1 no calculable sin tiempos de C0.")
        c0_results = []

    # 3. Ejecutar C1, C2, C3
    print("\n[3/6] C1 — LLM puro...")
    c1_results = run_c1_llm_pure(eval_pairs)

    print("\n[4/6] C2 — LLM + RAG...")
    c2_results = run_c2_llm_rag(eval_pairs)

    print("\n[5/6] C3 — LLM + RAG + PII...")
    c3_results = run_c3_llm_rag_pii(eval_pairs)

    # 4. Calcular métricas
    print("\n[6/6] Calculando métricas H1 / H2 / H3...")

    # H2 — Eficacia (C1, C2, C3 vs Gold Standard)
    h2_c1 = compute_h2_metrics(c1_results, "C1")
    h2_c2 = compute_h2_metrics(c2_results, "C2")
    h2_c3 = compute_h2_metrics(c3_results, "C3")

    # H3 — Equidad (C2 vs C3)
    h3_c2 = compute_h3_metrics(c2_results, "C2")
    h3_c3 = compute_h3_metrics(c3_results, "C3")

    # H1 — Eficiencia (C0 vs C1/C2/C3)
    h1_c1 = compute_h1_metrics(c0_results, c1_results, "C1") if c0_results else {}
    h1_c2 = compute_h1_metrics(c0_results, c2_results, "C2") if c0_results else {}
    h1_c3 = compute_h1_metrics(c0_results, c3_results, "C3") if c0_results else {}

    # 4.5. Guardar tablas H1
    h1_headers = ["Configuración", "Mediana C0 (s)", "Mediana Cx (s)", "IQR C0 (s)", "IQR Cx (s)", "Factor Speedup", "U Mann-Whitney", "p-valor", "H1 aceptada"]
    h1_rows = []
    if c0_results:
        for name, h1_res in [("C1 (LLM puro)", h1_c1), ("C2 (LLM + RAG)", h1_c2), ("C3 (LLM + RAG + PII)", h1_c3)]:
            if h1_res:
                h1_rows.append([
                    name,
                    f"{h1_res.get('median_baseline_s', 0.0):.1f}",
                    f"{h1_res.get('median_system_s', 0.0):.1f}",
                    f"{h1_res.get('IQR_baseline_s', 0.0):.1f}",
                    f"{h1_res.get('IQR_system_s', 0.0):.1f}",
                    f"{h1_res.get('speedup_factor', 0.0):.1f}x",
                    f"{h1_res.get('mannwhitney_U', 0.0):.1f}",
                    f"{h1_res.get('p_value', 0.0):.4f}",
                    "Sí" if h1_res.get("H1_accepted") else "No"
                ])
    else:
        h1_rows = [
            ["C1 (LLM puro)", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A", "No"],
            ["C2 (LLM + RAG)", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A", "No"],
            ["C3 (LLM + RAG + PII)", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A", "N/A", "No"],
        ]
    save_word_table(
        headers=h1_headers,
        rows=h1_rows,
        caption="Tabla 7.3. Métricas de eficiencia por configuración (H1)",
        filename="tab_resultados_h1.docx",
        note=(
            "Medianas e IQR expresados en segundos por candidato. "
            "U Mann-Whitney y p-valor corresponden a la comparación unilateral con C0."
        ),
    )
    save_csv(
        [{"config": r[0], "median_c0": r[1], "median_cx": r[2], "speedup": r[5], "p_value": r[7], "h1_accepted": r[8]} for r in h1_rows],
        "tab_resultados_h1.csv",
    )

    # 5. Guardar tablas H2
    h2_headers = ["Configuración", "F1-score (macro)", "AUC-ROC", "IC 95% AUC-ROC", "H2 aceptada"]
    h2_rows = [
        ["C1 (LLM puro)",
         f"{h2_c1.get('f1', 'N/A'):.3f}" if h2_c1 else "N/A",
         f"{h2_c1.get('auc_roc', 'N/A'):.3f}" if h2_c1 else "N/A",
         f"({h2_c1.get('auc_ci_low', 'N/A'):.3f}, {h2_c1.get('auc_ci_high', 'N/A'):.3f})" if h2_c1 else "N/A",
         "Sí" if h2_c1 and h2_c1.get("f1", 0) >= 0.85 and h2_c1.get("auc_roc", 0) >= 0.90 else "No"],
        ["C2 (LLM + RAG)",
         f"{h2_c2.get('f1', 'N/A'):.3f}" if h2_c2 else "N/A",
         f"{h2_c2.get('auc_roc', 'N/A'):.3f}" if h2_c2 else "N/A",
         f"({h2_c2.get('auc_ci_low', 'N/A'):.3f}, {h2_c2.get('auc_ci_high', 'N/A'):.3f})" if h2_c2 else "N/A",
         "Sí" if h2_c2 and h2_c2.get("f1", 0) >= 0.85 and h2_c2.get("auc_roc", 0) >= 0.90 else "No"],
        ["C3 (LLM + RAG + PII)",
         f"{h2_c3.get('f1', 'N/A'):.3f}" if h2_c3 else "N/A",
         f"{h2_c3.get('auc_roc', 'N/A'):.3f}" if h2_c3 else "N/A",
         f"({h2_c3.get('auc_ci_low', 'N/A'):.3f}, {h2_c3.get('auc_ci_high', 'N/A'):.3f})" if h2_c3 else "N/A",
         "Sí" if h2_c3 and h2_c3.get("f1", 0) >= 0.85 and h2_c3.get("auc_roc", 0) >= 0.90 else "No"],
    ]
    save_word_table(
        headers=h2_headers,
        rows=h2_rows,
        caption="Tabla 7.1. Métricas de eficacia por configuración frente al Gold Standard (H2)",
        filename="tab_resultados_h2.docx",
        note=(
            "IC 95% calculados con bootstrap (B = 1000). "
            "Umbral H2: F1 ≥ 0.85 y AUC-ROC ≥ 0.90. "
            "Gold Standard: κ ≥ 0.70 inter-evaluador."
        ),
    )
    save_csv(
        [{"config": r[0], "f1": r[1], "auc_roc": r[2]} for r in h2_rows],
        "tab_resultados_h2.csv",
    )

    # Guardar tablas H3
    h3_headers = ["Configuración", "DIR (género)", "SPD (género)", "H3 aceptada (DIR ≥ 0.80)"]
    h3_rows = [
        ["C2 (LLM + RAG)",
         f"{h3_c2.get('dir', 'N/A'):.3f}" if h3_c2 else "N/A",
         f"{h3_c2.get('spd', 'N/A'):.3f}" if h3_c2 else "N/A",
         "Sí" if h3_c2 and h3_c2.get("dir", 0) >= 0.80 else "No"],
        ["C3 (LLM + RAG + PII)",
         f"{h3_c3.get('dir', 'N/A'):.3f}" if h3_c3 else "N/A",
         f"{h3_c3.get('spd', 'N/A'):.3f}" if h3_c3 else "N/A",
         "Sí" if h3_c3 and h3_c3.get("dir", 0) >= 0.80 else "No"],
    ]
    save_word_table(
        headers=h3_headers,
        rows=h3_rows,
        caption="Tabla 7.2. Métricas de equidad algorítmica por configuración (H3)",
        filename="tab_resultados_h3.docx",
        note=(
            "DIR = Disparate Impact Ratio (umbral EEOC: ≥ 0.80). "
            "SPD = Statistical Parity Difference (ideal: 0). "
            "Grupo protegido: género (M/F). "
            "C2 sin anonimización vs C3 con anonimización PII."
        ),
    )
    save_csv(
        [{"config": r[0], "dir": r[1], "spd": r[2]} for r in h3_rows],
        "tab_resultados_h3.csv",
    )

    try:
        from evaluation.export_excel_report import generate_excel_report
        generate_excel_report()
    except Exception as e:
        print(f"  [WARN] No se pudo generar el reporte Excel: {e}")

    print("\n" + "=" * 60)
    print("Experimento completado. Outputs en paper/tables/")
    print("Insertar tablas en SISTAC_TFE.docx sección 7.")
    print("=" * 60)



def _generate_template_tables() -> None:
    """
    Genera tablas template (sin datos reales) para verificar el pipeline
    de exportación antes de tener el corpus completo.
    """
    print("\nGenerando tablas template (sin datos reales)...")

    save_word_table(
        headers=["Configuración", "F1-score (macro)", "AUC-ROC", "IC 95% AUC-ROC", "H2 aceptada"],
        rows=[
            ["C1 (LLM puro)",        "TODO", "TODO", "(TODO, TODO)", "TODO"],
            ["C2 (LLM + RAG)",       "TODO", "TODO", "(TODO, TODO)", "TODO"],
            ["C3 (LLM + RAG + PII)", "TODO", "TODO", "(TODO, TODO)", "TODO"],
        ],
        caption="Tabla 7.1. Métricas de eficacia por configuración frente al Gold Standard (H2) [TEMPLATE]",
        filename="tab_resultados_h2_template.docx",
        note="TODO = completar tras experimentos. Umbral H2: F1 ≥ 0.85, AUC-ROC ≥ 0.90.",
    )

    save_word_table(
        headers=["Configuración", "DIR (género)", "SPD (género)", "H3 aceptada"],
        rows=[
            ["C2 (LLM + RAG)",       "TODO", "TODO", "TODO"],
            ["C3 (LLM + RAG + PII)", "TODO", "TODO", "TODO"],
        ],
        caption="Tabla 7.2. Métricas de equidad algorítmica (H3) [TEMPLATE]",
        filename="tab_resultados_h3_template.docx",
        note="TODO = completar tras experimentos. Umbral H3: DIR ≥ 0.80 (EEOC 4/5).",
    )

    print("Templates generados en paper/tables/")


if __name__ == "__main__":
    main()
