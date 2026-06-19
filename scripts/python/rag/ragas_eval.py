"""
rag/ragas_eval.py — Evaluación técnica del pipeline RAG (sección 5.6, H2)

Calcula métricas RAGAS para evaluar la calidad del pipeline C2 en modo in-vitro
(antes del experimento factorial completo):
  - faithfulness:       ¿la respuesta/justificación está respaldada por los chunks?
  - answer_relevancy:   ¿la justificación responde a la JD?
  - context_precision:  ¿los chunks recuperados son relevantes para la respuesta?

LLM juez: Claude Haiku (via langchain-anthropic) — no OpenAI por defecto.

Si RAGAS no está instalado, se ejecutan métricas proxy basadas en ROUGE-L y
similitud coseno, que son suficientes para el análisis descriptivo del TFE.

Output: paper/tables/tab_ragas_c2.csv + tab_ragas_c2.docx

Hipótesis: H2 — las configuraciones RAG (C2/C3) alcanzan F1 ≥ 0.85 y
           AUC-ROC ≥ 0.90 frente al Gold Standard (experimento factorial).
           Esta evaluación es in-vitro — mide calidad interna del pipeline,
           no las métricas de hipótesis definitivas.

Uso:
    python scripts/python/rag/ragas_eval.py
    # Lee results/pilot_c2_results.json como dataset de evaluación
    # Genera paper/tables/tab_ragas_c2.csv y .docx
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

# INV-16: rutas via PROJECT_ROOT
_SCRIPTS_DIR = Path(__file__).parent.parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

from config import FIGURES_DIR, TABLES_DIR, PROJECT_ROOT, ensure_dirs

# ── Constantes ────────────────────────────────────────────────────────────────

PILOT_RESULTS_PATH = PROJECT_ROOT / "results" / "pilot_c2_results.json"
OUTPUT_CSV  = TABLES_DIR / "tab_ragas_c2.csv"
OUTPUT_DOCX = TABLES_DIR / "tab_ragas_c2.docx"


# ── RAGAS evaluation ──────────────────────────────────────────────────────────

def run_ragas_evaluation(eval_dataset: list[dict]) -> dict:
    """
    Evalúa el pipeline RAG con métricas RAGAS usando Claude como LLM juez.

    Args:
        eval_dataset: Lista de diccionarios con claves:
            - question:     La JD usada como consulta de retrieval
            - answer:       La justificación generada por el LLM
            - contexts:     Lista de chunks recuperados por Azure Search
            - ground_truth: Evaluación del experto RRHH (texto o etiqueta)

    Returns:
        Diccionario con métricas promedio:
            {faithfulness, answer_relevancy, context_precision}

    Raises:
        ImportError: si ragas o langchain-anthropic no están instalados.
    """
    try:
        from ragas import evaluate
        from ragas.metrics import faithfulness, answer_relevancy, context_precision
        from langchain_anthropic import ChatAnthropic
        from datasets import Dataset
    except ImportError as exc:
        raise ImportError(
            "RAGAS o langchain-anthropic no instalados.\n"
            "Instalar con: pip install ragas>=0.1.0 langchain-anthropic>=0.2.0\n"
            "Alternativa: usar run_proxy_evaluation() para métricas ROUGE/coseno."
        ) from exc

    llm_judge = ChatAnthropic(model="claude-haiku-4-5-20241022")

    dataset = Dataset.from_list(eval_dataset)

    result = evaluate(
        dataset=dataset,
        metrics=[faithfulness, answer_relevancy, context_precision],
        llm=llm_judge,
    )

    return {
        "faithfulness":      round(result["faithfulness"], 4),
        "answer_relevancy":  round(result["answer_relevancy"], 4),
        "context_precision": round(result["context_precision"], 4),
        "n_samples":         len(eval_dataset),
        "method":            "ragas_claude",
    }


def _rouge_l(candidate: str, reference: str) -> float:
    """
    Calcula el score ROUGE-L aproximado (LCS F-measure) entre dos textos.
    """
    if not candidate or not reference:
        return 0.0

    cand_tokens = [w.strip(".,;:?!()[]{}'\"").lower() for w in candidate.split() if w.strip()]
    ref_tokens = [w.strip(".,;:?!()[]{}'\"").lower() for w in reference.split() if w.strip()]

    if not cand_tokens or not ref_tokens:
        return 0.0

    m = len(cand_tokens)
    n = len(ref_tokens)
    dp = [0] * (n + 1)

    for i in range(1, m + 1):
        prev = 0
        for j in range(1, n + 1):
            temp = dp[j]
            if cand_tokens[i - 1] == ref_tokens[j - 1]:
                dp[j] = prev + 1
            else:
                dp[j] = max(dp[j], dp[j - 1])
            prev = temp

    lcs_len = dp[n]

    precision = lcs_len / m
    recall = lcs_len / n

    if precision + recall == 0:
        return 0.0
    return (2 * precision * recall) / (precision + recall)


def run_proxy_evaluation(eval_dataset: list[dict]) -> dict:
    """
    Métricas proxy cuando RAGAS no está disponible.

    - Faithfulness:       ROUGE-L entre justificación y chunks concatenados.
    - Context Precision:  Similitud coseno promedio entre chunks y la JD (question).
    - Answer Relevancy:   Similitud coseno entre justificación y JD.

    Args:
        eval_dataset: misma estructura que run_ragas_evaluation().

    Returns:
        Diccionario con métricas promedio proxy.
    """
    try:
        from sklearn.metrics.pairwise import cosine_similarity
        from sentence_transformers import SentenceTransformer
    except ImportError as exc:
        raise ImportError(
            "scikit-learn o sentence-transformers no instalados.\n"
            "pip install scikit-learn sentence-transformers"
        ) from exc

    model = SentenceTransformer("paraphrase-multilingual-mpnet-base-v2")

    faithfulness_scores = []
    context_precision_scores = []
    answer_relevancy_scores = []

    for item in eval_dataset:
        question   = item.get("question", "")
        answer     = item.get("answer", "")
        contexts   = item.get("contexts", [])

        # Faithfulness: ROUGE-L entre answer y contexto concatenado
        context_concat = " ".join(contexts)
        faithfulness_scores.append(_rouge_l(answer, context_concat))

        # Context Precision: similitud coseno chunks vs question
        if contexts:
            q_emb = model.encode([question])
            c_embs = model.encode(contexts)
            sims = cosine_similarity(q_emb, c_embs)[0]
            context_precision_scores.append(float(sims.mean()))
        else:
            context_precision_scores.append(0.0)

        # Answer Relevancy: similitud coseno answer vs question
        q_emb   = model.encode([question])
        a_emb   = model.encode([answer])
        sim     = cosine_similarity(q_emb, a_emb)[0][0]
        answer_relevancy_scores.append(float(sim))

    return {
        "faithfulness":      round(sum(faithfulness_scores) / len(faithfulness_scores), 4),
        "answer_relevancy":  round(sum(answer_relevancy_scores) / len(answer_relevancy_scores), 4),
        "context_precision": round(sum(context_precision_scores) / len(context_precision_scores), 4),
        "n_samples":         len(eval_dataset),
        "method":            "proxy_rouge_cosine",
    }


# ── Carga de datos del piloto ─────────────────────────────────────────────────

def load_pilot_as_ragas_dataset(path: Path = PILOT_RESULTS_PATH) -> list[dict]:
    """
    Convierte pilot_c2_results.json al formato esperado por RAGAS.

    El piloto tiene: cv_id, jd_id, score, justification, decision,
    expected_label, expected_score.
    RAGAS espera: question, answer, contexts, ground_truth.

    Mapeo:
        question     ← jd_id (usamos el ID; la JD completa requiere data/raw/jobs/)
        answer       ← justification
        contexts     ← [] (el piloto no guardó los chunks recuperados)
        ground_truth ← expected_label

    Nota: contexts vacío implica que context_precision no es evaluable con el
    piloto. Para RAGAS completo, correr rag/pipeline.py evaluate() que sí
    guarda los chunks.
    """
    if not path.exists():
        raise FileNotFoundError(
            f"No se encontró el archivo de resultados del piloto: {path}\n"
            "Correr evaluate_pilot_c2.py primero."
        )

    with open(path, encoding="utf-8") as f:
        pilot_data = json.load(f)

    ragas_dataset = []
    for item in pilot_data:
        ragas_dataset.append({
            "question":     f"Cargo: {item.get('job_id', '')}",
            "answer":       item.get("justification", ""),
            "contexts":     [],  # piloto no guardó chunks — limitación documentada
            "ground_truth": item.get("expected_label", ""),
        })

    return ragas_dataset


# ── Guardar resultados ────────────────────────────────────────────────────────

def save_ragas_results(metrics: dict) -> None:
    """Guarda métricas RAGAS en CSV y DOCX para el paper."""
    import csv

    ensure_dirs()

    # CSV
    with open(OUTPUT_CSV, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=metrics.keys())
        writer.writeheader()
        writer.writerow(metrics)
    print(f"  → CSV: {OUTPUT_CSV}")

    # DOCX
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        p = doc.add_paragraph()
        run = p.add_run(
            "Tabla 5.1. Métricas de evaluación técnica del pipeline RAG (in-vitro, C2)"
        )
        run.bold = True
        run.font.size = Pt(11)

        headers = ["Métrica", "Valor", "Interpretación", "Método"]
        interpretaciones = {
            "faithfulness":      "Respuestas respaldadas por contexto",
            "answer_relevancy":  "Relevancia de la justificación vs JD",
            "context_precision": "Proporción de chunks útiles recuperados",
        }

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        for metric in ["faithfulness", "answer_relevancy", "context_precision"]:
            if metric in metrics:
                row = table.add_row().cells
                row[0].text = metric.replace("_", " ").title()
                row[1].text = str(metrics[metric])
                row[2].text = interpretaciones.get(metric, "")
                row[3].text = metrics.get("method", "")

        p_note = doc.add_paragraph()
        run_note = p_note.add_run(
            f"Nota. n = {metrics.get('n_samples', '?')} pares CV-JD (piloto). "
            f"Método: {metrics.get('method', '?')}. "
            "Los valores oscilan entre 0 y 1; valores más cercanos a 1 indican "
            "mayor calidad del pipeline. Context precision no evaluable en el "
            "piloto (chunks no almacenados); evaluación completa post-corpus."
        )
        run_note.italic = True
        run_note.font.size = Pt(10)

        doc.save(str(OUTPUT_DOCX))
        print(f"  → DOCX: {OUTPUT_DOCX}")

    except ImportError:
        print("  [WARN] python-docx no instalado — solo se generó CSV.")


# ── Main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=== SISTAC — Evaluación RAGAS del Pipeline C2 ===\n")

    # Cargar dataset del piloto
    try:
        dataset = load_pilot_as_ragas_dataset()
        print(f"Dataset cargado: {len(dataset)} muestras del piloto C2")
    except FileNotFoundError as e:
        print(f"[ERROR] {e}")
        sys.exit(1)

    # Intentar RAGAS primero; fallback a proxy
    print("\nIntentando evaluación RAGAS (Claude como juez)...")
    try:
        metrics = run_ragas_evaluation(dataset)
        print("  ✓ RAGAS con Claude Haiku")
    except ImportError as e:
        print(f"  [WARN] {e}")
        print("\nFallback: evaluación proxy (ROUGE-L + coseno)...")
        metrics = run_proxy_evaluation(dataset)
        print(f"  ✓ Proxy ({metrics['method']})")

    # Mostrar resultados
    print(f"\n{'Métrica':<25} {'Valor':>8}")
    print("-" * 35)
    for k in ["faithfulness", "answer_relevancy", "context_precision"]:
        if k in metrics:
            print(f"  {k:<23} {metrics[k]:>8.4f}")
    print(f"\n  n = {metrics['n_samples']} muestras | método: {metrics['method']}")

    # Guardar
    print("\nGuardando resultados...")
    save_ragas_results(metrics)
    print("\nListo. Insertar tab_ragas_c2.docx en SISTAC_TFE.docx sección 5.6.")
