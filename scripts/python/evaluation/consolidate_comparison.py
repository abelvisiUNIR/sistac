"""
scripts/python/evaluation/consolidate_comparison.py
Consolida los resultados de los diferentes LLM (Anthropic Claude, Google Gemini, OpenAI GPT)
en una única tabla comparativa para el Capítulo 5 de la tesis.
"""

from __future__ import annotations

import csv
from pathlib import Path
import sys

# INV-16: rutas via PROJECT_ROOT
_SCRIPTS_DIR = Path(__file__).parent.parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

from config import PROJECT_ROOT, TABLES_DIR


def load_provider_metrics(provider: str) -> dict | None:
    provider_dir = TABLES_DIR / provider
    h1_path = provider_dir / "tab_resultados_h1.csv"
    h2_path = provider_dir / "tab_resultados_h2.csv"
    h3_path = provider_dir / "tab_resultados_h3.csv"

    if not (h1_path.exists() and h2_path.exists() and h3_path.exists()):
        return None

    metrics = {}

    # Cargar H1 (Eficiencia)
    with open(h1_path, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            cfg = row["config"].lower()
            key_suffix = "c1" if "c1" in cfg else ("c2" if "c2" in cfg else "c3")
            metrics[f"t_med_{key_suffix}"] = f"{float(row['median_cx']):.1f}s"
            metrics[f"speedup_{key_suffix}"] = row["speedup"]

    # Cargar H2 (Eficacia)
    with open(h2_path, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            cfg = row["config"].lower()
            key_suffix = "c1" if "c1" in cfg else ("c2" if "c2" in cfg else "c3")
            metrics[f"f1_{key_suffix}"] = f"{float(row['f1']):.3f}"
            metrics[f"auc_{key_suffix}"] = f"{float(row['auc_roc']):.3f}"

    # Cargar H3 (Equidad)
    with open(h3_path, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            cfg = row["config"].lower()
            key_suffix = "c2" if "c2" in cfg else "c3"
            metrics[f"dir_{key_suffix}"] = f"{float(row['dir']):.3f}"
            metrics[f"spd_{key_suffix}"] = f"{float(row['spd']):.3f}"

    return metrics


def main():
    providers = ["anthropic", "google", "openai"]
    provider_names = {
        "anthropic": "Claude 4.5 Sonnet (Anthropic)",
        "google": "Gemini 2.5 Flash (Google)",
        "openai": "GPT-4o / mini (OpenAI)"
    }

    loaded_data = {}
    for p in providers:
        data = load_provider_metrics(p)
        if data:
            loaded_data[p] = data
            print(f"[Consolidador] Cargadas m\u00e9tricas para {p.upper()}")
        else:
            print(f"[Consolidador] Sin datos completos para {p.upper()} - omitiendo de la comparativa.")

    if not loaded_data:
        print("[ERROR] No hay datos de ning\u00fan proveedor para consolidar.")
        sys.exit(1)

    # Definir filas de la tabla
    row_definitions = [
        # Secci\u00f3n, M\u00e9trica, Clave M\u00e9trica
        ("Eficiencia", "Mediana T_cand (C1 - LLM Puro)", "t_med_c1"),
        ("Eficiencia", "Mediana T_cand (C2 - LLM + RAG)", "t_med_c2"),
        ("Eficiencia", "Mediana T_cand (C3 - RAG + PII)", "t_med_c3"),
        ("Eficiencia", "Factor Speedup (C1 - LLM Puro)", "speedup_c1"),
        ("Eficiencia", "Factor Speedup (C2 - LLM + RAG)", "speedup_c2"),
        ("Eficiencia", "Factor Speedup (C3 - RAG + PII)", "speedup_c3"),
        
        ("Eficacia", "F1-score macro (C1 - LLM Puro)", "f1_c1"),
        ("Eficacia", "F1-score macro (C2 - LLM + RAG)", "f1_c2"),
        ("Eficacia", "F1-score macro (C3 - RAG + PII)", "f1_c3"),
        ("Eficacia", "AUC-ROC (C1 - LLM Puro)", "auc_c1"),
        ("Eficacia", "AUC-ROC (C2 - LLM + RAG)", "auc_c2"),
        ("Eficacia", "AUC-ROC (C3 - RAG + PII)", "auc_c3"),
        
        ("Equidad (G\u00e9nero)", "Disparate Impact Ratio - DIR (C2)", "dir_c2"),
        ("Equidad (G\u00e9nero)", "Disparate Impact Ratio - DIR (C3)", "dir_c3"),
        ("Equidad (G\u00e9nero)", "Statistical Parity Diff - SPD (C2)", "spd_c2"),
        ("Equidad (G\u00e9nero)", "Statistical Parity Diff - SPD (C3)", "spd_c3"),
    ]

    # Construir cabeceras
    active_providers = list(loaded_data.keys())
    headers = ["Categor\u00eda", "M\u00e9trica / Configuraci\u00f3n"] + [provider_names[p] for p in active_providers]

    # Construir filas de datos
    table_rows = []
    for cat, metric_label, key in row_definitions:
        row = [cat, metric_label]
        for p in active_providers:
            val = loaded_data[p].get(key, "N/A")
            row.append(val)
        table_rows.append(row)

    # 1. Guardar como CSV
    csv_path = TABLES_DIR / "tab_comparativa_modelos.csv"
    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        writer.writerows(table_rows)
    print(f"[OK] CSV de comparaci\u00f3n creado: {csv_path.name}")

    # 2. Guardar como Markdown
    md_path = TABLES_DIR / "comparativa_modelos.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("# Comparativa General de Desempe\u00f1o por Modelo LLM\n\n")
        f.write("Esta tabla compara el desempe\u00f1o de los distintos modelos evaluados para las hip\u00f3tesis H1 (Eficiencia), H2 (Eficacia) y H3 (Equidad).\n\n")
        
        # Escribir cabecera Markdown
        f.write("| " + " | ".join(headers) + " |\n")
        f.write("|" + "|".join(["---"] * len(headers)) + "|\n")
        
        # Escribir filas
        for row in table_rows:
            f.write("| " + " | ".join(row) + " |\n")
    print(f"[OK] Markdown de comparaci\u00f3n creado: {md_path.name}")

    # 3. Guardar como Word (.docx)
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        p = doc.add_paragraph()
        run = p.add_run("Tabla comparativa integrada de modelos (Claude vs. Gemini vs. GPT)")
        run.bold = True
        run.font.size = Pt(11)
        
        table = doc.add_table(rows=1 + len(table_rows), cols=len(headers))
        table.style = "Table Grid"
        
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        for r_idx, row in enumerate(table_rows):
            row_cells = table.rows[r_idx + 1].cells
            for c_idx, val in enumerate(row):
                row_cells[c_idx].text = str(val)
                row_cells[c_idx].paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if c_idx < 2 else WD_ALIGN_PARAGRAPH.CENTER
                )
                
        docx_path = TABLES_DIR / "tab_comparativa_modelos.docx"
        doc.save(str(docx_path))
        print(f"[OK] DOCX de comparaci\u00f3n creado: {docx_path.name}")
    except ImportError:
        print("[WARN] python-docx no instalado. Omitiendo generaci\u00f3n de tabla Word.")


if __name__ == "__main__":
    main()
