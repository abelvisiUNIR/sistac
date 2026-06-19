"""
scripts/python/utils/insert_google_migration_docx.py
Lee el archivo markdown paper/sections/estudio_datos_y_migracion_google.md y lo inserta
al final del documento paper/SISTAC_TFE.docx de manera limpia y sin tocar
el contenido actual.
"""

from __future__ import annotations

import shutil
import sys
from datetime import datetime
from pathlib import Path

# Asegurar codificación utf-8
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

_PROJECT_ROOT = Path(__file__).parent.parent.parent.parent
DOCX_PATH     = _PROJECT_ROOT / "paper" / "SISTAC_TFE.docx"
MD_PATH       = _PROJECT_ROOT / "paper" / "sections" / "estudio_datos_y_migracion_google.md"
BACKUP_DIR    = _PROJECT_ROOT / "paper" / "backups"

BACKUP_DIR.mkdir(parents=True, exist_ok=True)

# 1. Crear copia de seguridad
backup_name = f"SISTAC_TFE_backup_gcp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
shutil.copy2(DOCX_PATH, BACKUP_DIR / backup_name)
print(f"[INFO] Backup creado: {backup_name}")

# 2. Cargar dependencias de python-docx
try:
    import docx
except ImportError:
    print("[ERROR] Requiere python-docx. Instalando...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

# 3. Leer y parsear el archivo Markdown
if not MD_PATH.exists():
    print(f"[ERROR] No se encontró el archivo markdown en {MD_PATH}")
    sys.exit(1)

print(f"[INFO] Leyendo borrador desde {MD_PATH.name}...")
md_content = MD_PATH.read_text(encoding="utf-8")
lines = md_content.splitlines()

# 4. Cargar el documento de Word
print(f"[INFO] Cargando documento {DOCX_PATH.name}...")
doc = docx.Document(str(DOCX_PATH))

# 5. Insertar salto de sección / página antes de añadir el contenido
doc.add_page_break()
print("[INFO] Añadido salto de página para nueva sección...")

# 6. Parsear líneas simples de markdown e insertar en docx
in_code_block = False

for line in lines:
    stripped = line.strip()
    
    # Manejar bloques de código (saltar Mermaid o código formateado)
    if stripped.startswith("```"):
        in_code_block = not in_code_block
        continue
        
    if in_code_block:
        # En bloques de código, añadir texto como fuente monoespaciada
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = docx.shared.Inches(0.5)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = docx.shared.Pt(9.5)
        continue

    if not stripped:
        continue

    # Encabezados
    if stripped.startswith("# "):
        text = stripped[2:]
        doc.add_heading(text, level=1)
        print(f"  Añadido Título 1: {text}")
    elif stripped.startswith("## "):
        text = stripped[3:]
        doc.add_heading(text, level=2)
        print(f"  Añadido Título 2: {text}")
    elif stripped.startswith("### "):
        text = stripped[4:]
        doc.add_heading(text, level=3)
        print(f"  Añadido Título 3: {text}")
    # Listas
    elif stripped.startswith("* "):
        doc.add_paragraph("• " + stripped[2:])
    elif stripped.startswith("- "):
        doc.add_paragraph("• " + stripped[2:])
    # Notas / Citas
    elif stripped.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = docx.shared.Inches(0.5)
        run = p.add_run(stripped[2:])
        run.italic = True
    # Párrafo común
    else:
        # Reemplazar formatos simples como **negrita**
        p = doc.add_paragraph()
        parts = stripped.split("**")
        for idx, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            if idx % 2 == 1:
                run.bold = True

# 7. Guardar el documento
print(f"[INFO] Guardando cambios en {DOCX_PATH.name}...")
doc.save(str(DOCX_PATH))
print("[SUCCESS] El documento Word ha sido actualizado correctamente.")
