"""
docx_extractor.py — Extracción de texto plano desde archivos .docx

Extrae el contenido de los documentos fuente para facilitar la migración
a LaTeX (Paso 7 del plan de implementación).

Uso:
    python scripts/python/utils/docx_extractor.py

Outputs (en explorations/raw/):
    - ch1_intro_raw.txt
    - ch2_litreview_raw.txt
    - ch3_methodology_raw.txt
    - ch4_system_design_raw.txt
"""

import sys
from pathlib import Path

# Rutas relativas (INV-16: no rutas absolutas)
SCRIPTS_DIR = Path(__file__).parent.parent
PROJECT_ROOT = SCRIPTS_DIR.parent.parent
sys.path.insert(0, str(SCRIPTS_DIR))

from config import SUPPORTING_DOCS, EXPLORATIONS

try:
    from docx import Document
except ImportError:
    print("ERROR: Instalar python-docx con: pip install python-docx")
    sys.exit(1)


def extract_docx_text(docx_path: Path) -> str:
    """
    Extrae texto plano de un archivo .docx preservando estructura de párrafos.

    Args:
        docx_path: Ruta al archivo .docx

    Returns:
        Texto extraído como cadena, con párrafos separados por líneas en blanco.
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"No se encontró el archivo: {docx_path}")

    doc = Document(str(docx_path))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Detectar headings por estilo
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    prefix = "#" * int(level)
                except ValueError:
                    prefix = "#"
                paragraphs.append(f"\n{prefix} {text}\n")
            else:
                paragraphs.append(text)

    return "\n\n".join(paragraphs)


def extract_tables(docx_path: Path) -> str:
    """
    Extrae tablas de un .docx como texto delimitado por pipes.

    Args:
        docx_path: Ruta al archivo .docx

    Returns:
        Representación de texto de las tablas encontradas.
    """
    doc = Document(str(docx_path))
    tables_text = []

    for i, table in enumerate(doc.tables, 1):
        tables_text.append(f"\n--- TABLA {i} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            tables_text.append(" | ".join(cells))

    return "\n".join(tables_text)


def main():
    """Extrae todos los documentos fuente a texto plano."""
    output_dir = EXPLORATIONS / "raw"
    output_dir.mkdir(parents=True, exist_ok=True)

    sources = {
        "main_thesis": (
            SUPPORTING_DOCS / "Optimizacion_del_Proceso_de_Seleccion_de_Talento_.docx",
            "main_thesis_full_raw.txt",
        ),
        "lit_review": (
            SUPPORTING_DOCS / "SISTAC_Revision_Literatura.docx",
            "lit_review_raw.txt",
        ),
    }

    for name, (src_path, out_filename) in sources.items():
        out_path = output_dir / out_filename
        try:
            print(f"Extrayendo {src_path.name}...")
            text = extract_docx_text(src_path)
            tables = extract_tables(src_path)

            with open(out_path, "w", encoding="utf-8") as f:
                f.write(f"# EXTRACCIÓN: {src_path.name}\n")
                f.write(f"# Generado por: utils/docx_extractor.py\n\n")
                f.write(text)
                if tables.strip():
                    f.write("\n\n" + "=" * 60 + "\nTABLAS\n" + "=" * 60)
                    f.write(tables)

            print(f"  → {out_path} ({len(text)} caracteres)")
        except FileNotFoundError as e:
            print(f"  ADVERTENCIA: {e}")
            print(f"  Copia el .docx a: {SUPPORTING_DOCS}/")

    print("\nListo. Archivos en:", output_dir)
    print("Siguiente paso: usar /write para migrar cada capítulo a LaTeX.")


if __name__ == "__main__":
    main()
