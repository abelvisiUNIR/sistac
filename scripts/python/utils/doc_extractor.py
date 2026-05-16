"""
utils/doc_extractor.py — Extracción de texto desde PDF, DOCX e imágenes

Soporta:
  - .pdf   → pdfplumber (texto nativo) o Claude Vision (PDF escaneado)
  - .docx  → python-docx
  - .doc   → advertencia: convertir a .docx primero
  - .png / .jpg / .jpeg / .webp / .gif → Claude Vision API (sin Tesseract)
  - .txt   → lectura directa UTF-8

Uso:
    from utils.doc_extractor import extract_text
    texto = extract_text(Path("cv_maria.pdf"))
    texto = extract_text(Path("cargo_analista.docx"))
    texto = extract_text(uploaded_bytes, filename="cv_scan.png")
"""

from __future__ import annotations

import base64
import sys
from pathlib import Path

# INV-16: rutas via PROJECT_ROOT
_SCRIPTS_DIR = Path(__file__).parent.parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".png", ".jpg", ".jpeg", ".webp", ".gif"}
IMAGE_EXTENSIONS     = {".png", ".jpg", ".jpeg", ".webp", ".gif"}


# ── Función principal ─────────────────────────────────────────────────────────

def extract_text(source: Path | bytes, filename: str = "") -> str:
    """
    Extrae texto legible de un documento o imagen.

    Args:
        source:   Path al archivo, o bytes del contenido (para uploads web).
        filename: Nombre del archivo (requerido si source es bytes).

    Returns:
        Texto extraído como string.

    Raises:
        ValueError: si el formato no está soportado.
        RuntimeError: si la extracción falla.
    """
    if isinstance(source, Path):
        filename = source.name
        data = source.read_bytes()
    else:
        data = source

    ext = Path(filename).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Formato '{ext}' no soportado.\n"
            f"Formatos aceptados: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if ext == ".txt":
        return _extract_txt(data)
    elif ext == ".docx":
        return _extract_docx(data)
    elif ext == ".doc":
        raise ValueError(
            "Formato .doc no soportado directamente.\n"
            "Convertí el archivo a .docx en Word: Archivo → Guardar como → .docx"
        )
    elif ext == ".pdf":
        return _extract_pdf(data, filename)
    elif ext in IMAGE_EXTENSIONS:
        return _extract_image_vision(data, filename)
    else:
        raise ValueError(f"Formato '{ext}' no implementado.")


# ── Extractores ───────────────────────────────────────────────────────────────

def _extract_txt(data: bytes) -> str:
    """Texto plano — intenta UTF-8, fallback a latin-1."""
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1")


def _extract_docx(data: bytes) -> str:
    """Extrae texto de un archivo .docx usando python-docx."""
    try:
        from docx import Document
        import io
    except ImportError:
        raise RuntimeError(
            "python-docx no instalado.\n"
            "pip install python-docx"
        )

    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # También extraer tablas (útil para CVs con formato de tabla)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    text = "\n".join(paragraphs)
    if len(text.strip()) < 20:
        raise RuntimeError(
            "El DOCX parece estar vacío o solo contiene imágenes.\n"
            "Probá exportarlo como PDF y subir el PDF."
        )
    return text


def _extract_pdf(data: bytes, filename: str = "documento.pdf") -> str:
    """
    Extrae texto de un PDF.

    Estrategia:
    1. pdfplumber para PDFs con texto nativo (la mayoría de CVs digitales)
    2. Si el texto extraído es muy corto (PDF escaneado), usa Claude Vision
    """
    try:
        import pdfplumber
        import io
    except ImportError:
        raise RuntimeError(
            "pdfplumber no instalado.\n"
            "pip install pdfplumber"
        )

    text_parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)

    text = "\n".join(text_parts).strip()

    # Si hay poco texto (PDF escaneado / imagen), usar Vision API
    if len(text) < 100:
        print(f"  [doc_extractor] PDF con poco texto ({len(text)} chars) — usando Claude Vision")
        return _extract_pdf_vision(data, filename)

    return text


def _extract_pdf_vision(data: bytes, filename: str) -> str:
    """
    Extrae texto de un PDF escaneado convirtiendo páginas a imagen y
    enviándolas a Claude Vision.
    """
    try:
        import pdfplumber
        import io
    except ImportError:
        raise RuntimeError("pdfplumber no instalado. pip install pdfplumber")

    # Extraer cada página como imagen via pdfplumber
    all_text = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            img = page.to_image(resolution=150)
            # Guardar en buffer PNG
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            buf.seek(0)
            page_bytes = buf.read()

            page_text = _extract_image_vision(page_bytes, f"{filename}_p{i}.png")
            all_text.append(page_text)

    return "\n\n".join(all_text)


def _extract_image_vision(data: bytes, filename: str) -> str:
    """
    Extrae texto de una imagen usando Claude Vision API.
    No requiere Tesseract OCR — usa el LLM directamente.
    """
    import anthropic
    import os

    api_key = os.getenv("ANTHROPIC_API_KEY", "")
    if not api_key:
        raise RuntimeError(
            "ANTHROPIC_API_KEY no configurada.\n"
            "Necesaria para procesar imágenes con Claude Vision."
        )

    ext = Path(filename).suffix.lower()
    media_type_map = {
        ".png":  "image/png",
        ".jpg":  "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif":  "image/gif",
    }
    media_type = media_type_map.get(ext, "image/png")

    image_data = base64.standard_b64encode(data).decode("utf-8")

    client = anthropic.Anthropic(api_key=api_key)
    response = client.messages.create(
        model="claude-haiku-4-5-20241022",
        max_tokens=2048,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": image_data,
                        },
                    },
                    {
                        "type": "text",
                        "text": (
                            "Extraé todo el texto de esta imagen de forma fiel y completa. "
                            "Si es un CV, preservá la estructura: nombre, contacto, experiencia, "
                            "educación, habilidades. Si es una descripción de cargo, preservá "
                            "título, responsabilidades y requisitos. "
                            "Solo devolvé el texto extraído, sin comentarios adicionales."
                        ),
                    },
                ],
            }
        ],
    )

    return response.content[0].text


# ── Utilidad: detectar tipo de contenido ──────────────────────────────────────

def detect_media_type(filename: str) -> str:
    """Retorna el media type MIME según la extensión."""
    ext = Path(filename).suffix.lower()
    types = {
        ".pdf":  "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt":  "text/plain",
        ".png":  "image/png",
        ".jpg":  "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
    }
    return types.get(ext, "application/octet-stream")


# ── Demo ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    parser = argparse.ArgumentParser(description="Extraer texto de un documento")
    parser.add_argument("archivo", help="Ruta al archivo (PDF, DOCX, imagen, TXT)")
    args = parser.parse_args()

    path = Path(args.archivo)
    if not path.exists():
        print(f"Archivo no encontrado: {path}")
        sys.exit(1)

    print(f"Extrayendo texto de: {path.name}")
    texto = extract_text(path)
    print(f"Caracteres extraídos: {len(texto)}")
    print("-" * 60)
    print(texto[:1000])
    if len(texto) > 1000:
        print(f"\n... ({len(texto) - 1000} caracteres más)")
