import sys
import docx

def extract_docx_to_md(docx_path, output_path):
    doc = docx.Document(docx_path)
    with open(output_path, "w", encoding="utf-8") as f:
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            
            # Identify headings based on style
            style_name = p.style.name.lower()
            if "heading 1" in style_name:
                f.write(f"\n# {text}\n\n")
            elif "heading 2" in style_name:
                f.write(f"\n## {text}\n\n")
            elif "heading 3" in style_name:
                f.write(f"\n### {text}\n\n")
            elif "heading 4" in style_name:
                f.write(f"\n#### {text}\n\n")
            elif "heading" in style_name:
                f.write(f"\n# {text}\n\n")
            else:
                f.write(f"{text}\n\n")
                
        # Also let's extract tables and print them in text format
        if doc.tables:
            f.write("\n\n# TABLAS DEL DOCUMENTO\n\n")
            for i, table in enumerate(doc.tables):
                f.write(f"\n## Tabla {i+1}\n\n")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip().replace('\n', ' '))
                    # Simple markdown-like representation
                    f.write("| " + " | ".join(row_text) + " |\n")
                f.write("\n")

if __name__ == "__main__":
    docx_file = r"C:\Users\abelvisi\Documents\Google_Drive\Mi unidad\Máster UNIR\IA Y Data\TFE\Entregas_TFE_Terminal\clo-author\paper\Talento sin nombre anonimización, LLMs y RAG en el cribado curricular grupo 4 ESIT.docx"
    output_file = r"C:\Users\abelvisi\Documents\Google_Drive\Mi unidad\Máster UNIR\IA Y Data\TFE\Entregas_TFE_Terminal\clo-author\paper\Talento_sin_nombre_extracted.md"
    
    print(f"Extracting {docx_file} ...")
    extract_docx_to_md(docx_file, output_file)
    print(f"Extraction completed. Saved to {output_file}")
